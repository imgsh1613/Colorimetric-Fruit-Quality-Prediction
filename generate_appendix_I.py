"""
Generate Appendix I — DATASET FEATURE DESCRIPTION
for FruitSense Minor Project Report.

Saves: FruitSense_Appendix_I.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "FruitSense_Appendix_I.docx")

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    """Thin black border on all sides."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))   # 1 cm ≈ 567 twips
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# ── Document setup ─────────────────────────────────────────────────────────────

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.0)

# ── APPENDIX LABEL ─────────────────────────────────────────────────────────────
ap = doc.add_paragraph("APPENDIX I")
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap.runs[0].bold      = True
ap.runs[0].font.size = Pt(13)
ap.paragraph_format.space_after = Pt(2)

# ── TITLE ──────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph("DATASET FEATURE DESCRIPTION")
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.runs[0].bold      = True
tp.runs[0].font.size = Pt(14)
tp.paragraph_format.space_after = Pt(14)

# ── INTRO PARAGRAPH ────────────────────────────────────────────────────────────
intro = doc.add_paragraph(
    "The following table presents information about every feature used in the "
    "FruitSense dataset by showing its data type and value range together with "
    "a short explanation of how it contributes to fruit quality classification "
    "and freshness scoring. The dataset merges two publicly available sources: "
    "the Fruits-360 dataset (fresh specimens) and the Rotten Fruits dataset "
    "(Kaggle), producing a six-class annotated corpus of 11,949 images. "
    "In addition to raw image pixels, the inference pipeline extracts "
    "CIE L*a*b* colorimetric features from each image to support the "
    "four-factor freshness scoring engine."
)
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro.runs[0].font.size = Pt(11)
intro.paragraph_format.space_after = Pt(14)

# ── TABLE DATA ─────────────────────────────────────────────────────────────────
headers = ["Feature Name", "Data Type", "Range", "Description"]

rows = [
    # ── Dataset / Image-level features ──────────────────────────────────────
    ("image_path",       "String",      "N/A",
     "Absolute file-system path to the source image (JPG / PNG / JPEG)"),

    ("class_label",      "Categorical", "6 classes",
     "Ground-truth class: Apple, Apple_Rotten, Banana, Banana_Rotten, "
     "Orange, or Orange_Rotten"),

    ("split",            "Categorical", "train / val / test",
     "Dataset partition to which the image belongs (stratified 70/15/15 split)"),

    ("image_width",      "Integer",     "100 – 640 px",
     "Original width of the source image before resizing to 224×224"),

    ("image_height",     "Integer",     "100 – 480 px",
     "Original height of the source image before resizing to 224×224"),

    ("class_weight",     "Float",       "1.0 – 12.0",
     "Per-class loss weight (w = N_total / (C × N_c)) used during training "
     "to compensate for fresh/rotten class imbalance"),

    # ── CIE L*a*b* Colorimetric features ────────────────────────────────────
    ("mean_L_star",      "Float",       "0.0 – 100.0",
     "CIE L* mean lightness of the image; decreases as decay darkens tissue "
     "(contributes 0–35 pts to freshness score)"),

    ("mean_a_star",      "Float",       "-128.0 – 127.0",
     "CIE a* chromaticity (green–red axis); positive values indicate red "
     "pigmentation typical of fresh apples"),

    ("mean_b_star",      "Float",       "-128.0 – 127.0",
     "CIE b* chromaticity (blue–yellow axis); healthy carotenoid pigments "
     "produce high positive b*; Maillard browning destroys it "
     "(contributes 0–25 pts to freshness score)"),

    ("chroma_C_star",    "Float",       "0.0 – 180.0",
     "Chroma C* = √(a*² + b*²); measures overall colour saturation; "
     "diminishes as pigments degrade during senescence "
     "(contributes 0–30 pts to freshness score)"),

    ("hue_angle_H",      "Float",       "0.0 – 360.0°",
     "Hue angle H° = arctan2(b*, a*); shifts toward 40–65° (brown zone) "
     "as Maillard browning accumulates; used in hue decay penalty P_H"),

    ("hue_decay_penalty","Float",       "0.0 – 10.0",
     "P_H penalty term; fires only when H° is in the brown zone AND chroma "
     "C* is low (< 30), preventing false penalisation of vibrant orange fruit"),

    # ── Derived / Inference output features ─────────────────────────────────
    ("freshness_score",  "Float",       "0.0 – 100.0",
     "Composite freshness index: S = S_L + S_C + S_b − P_H; "
     "≥ 60 → Fresh, 40–59 → Marginal, < 40 → Rotten/Discard"),

    ("decay_percentage", "Float",       "0.0 – 100.0",
     "Decay percentage = 100 − freshness_score; direct inverse of the "
     "freshness score representing extent of deterioration"),

    ("predicted_class",  "Categorical", "6 classes",
     "MobileNetV2 classifier output (softmax argmax); one of the six "
     "fresh/rotten fruit classes"),

    ("confidence",       "Float",       "0.0 – 1.0",
     "Softmax probability of the predicted class; mean value 0.9559 across "
     "the 1,161-image test set"),

    ("is_fresh",         "Binary",      "0 or 1",
     "Freshness flag: 1 = Fresh (predicted class does not contain '_Rotten'), "
     "0 = Rotten; derived from predicted_class label"),

    ("shelf_life_days",  "Integer",     "0 – 5",
     "Estimated days of remaining edible life: 5 (score ≥ 60), "
     "1 (40 ≤ score < 60), 0 (score < 40)"),

    ("weight_loss_pct",  "Float",       "0.000 – 1.000",
     "Predicted weight-loss fraction from SE-CNN regressor; "
     "proxy for desiccation and cellular moisture loss"),

    ("hardness_N",       "Float",       "0.0 – 50.0 N",
     "Predicted mechanical hardness (Newtons) from SE-CNN regressor; "
     "fresh fruit exhibits higher hardness values than decayed specimens"),
]

# ── Build table ────────────────────────────────────────────────────────────────
doc.add_paragraph()   # spacing before table

table = doc.add_table(rows=1 + len(rows), cols=4)
table.style        = 'Table Grid'
table.alignment    = WD_TABLE_ALIGNMENT.CENTER

# Column widths (cm): Feature Name | Data Type | Range | Description
col_widths = [4.0, 2.5, 3.0, 7.5]

# Header row
hdr_row = table.rows[0]
hdr_texts = headers
HDR_BG = '1B3A6B'   # dark navy — same as reference image
for ci, (hdr_text, cw) in enumerate(zip(hdr_texts, col_widths)):
    cell = hdr_row.cells[ci]
    cell.text = ''
    set_cell_bg(cell, HDR_BG)
    set_cell_borders(cell)
    set_col_width(cell, cw)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(hdr_text)
    run.bold            = True
    run.font.size       = Pt(10)
    run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)

# Data rows
for ri, (feat, dtype, rng, desc) in enumerate(rows):
    row     = table.rows[ri + 1]
    values  = [feat, dtype, rng, desc]
    row_bg  = 'F2F4FB' if ri % 2 == 1 else 'FFFFFF'   # subtle alternating shade

    for ci, (val, cw) in enumerate(zip(values, col_widths)):
        cell = row.cells[ci]
        cell.text = ''
        set_cell_bg(cell, row_bg)
        set_cell_borders(cell)
        set_col_width(cell, cw)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        # Feature name column: bold
        run = para.add_run(val)
        run.font.size = Pt(10)
        if ci == 0:
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif ci in (1, 2):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── TABLE CAPTION ──────────────────────────────────────────────────────────────
doc.add_paragraph()
cap = doc.add_paragraph("Table A.1: FruitSense Dataset Feature Descriptions")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.runs[0].bold      = True
cap.runs[0].font.size = Pt(11)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
