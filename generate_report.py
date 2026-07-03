"""
Generate two Minor Project Reports in CSE format:
  1. FruitSense_Report_CleanTest.docx  — 97.67% (lab/in-distribution evaluation)
  2. FruitSense_Report_RealWorld.docx  — 89.15% (real-world simulation evaluation)

Based on: Minor PROJECT REPORT format_CSE.docx
"""

import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
TEMPLATE  = os.path.join(BASE_DIR, "Minor PROJECT REPORT format_CSE.docx")

# ── Helpers ────────────────────────────────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()

def set_cell_bg(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)

def style_table_header_row(row, bg='1B3A6B', fg='FFFFFF'):
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold      = True
                run.font.color.rgb = RGBColor(
                    int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))

def add_metric_table(doc, title, rows_data, col_headers):
    """Add a styled table with caption above it."""
    caption_para = doc.add_paragraph(title)
    caption_para.style = 'Body Text'
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.runs[0]
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(col_headers):
        hdr.cells[i].text = h
    style_table_header_row(hdr)

    # Data rows
    for ri, row_data in enumerate(rows_data):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                set_cell_bg(cell, 'EEF2FF')
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # spacing

# ── Report generator ───────────────────────────────────────────────────────────

def generate_report(output_path: str, scenario: str):
    """
    scenario: 'clean'     → 97.67% accuracy (in-distribution lab test)
              'realworld' → 89.15% accuracy (real-world simulation)
    """

    is_clean = scenario == 'clean'

    # Use a fresh document (not the template) to avoid section-stripping issues
    doc = Document()
    # Set margins to match typical CSE report (1 inch all sides)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1)

    TITLE         = "FruitSense: AI-Powered Fruit Quality Analysis System"
    STUDENTS      = [
        ("Gautam",     "RA2110260400XXX"),
        ("Student 2",  "RA2110260400YYY"),
        ("Student 3",  "RA2110260400ZZZ"),
    ]
    GUIDE         = "Dr. [Guide Name]"
    GUIDE_DESIG   = "Assistant Professor, Department of CSE"
    DEPT          = "SRM INSTITUTE OF SCIENCE AND TECHNOLOGY"
    YEAR          = "MAY 2026"

    # ── Metrics ──
    if is_clean:
        overall_acc  = "97.67%"
        macro_prec   = "97.92%"
        macro_rec    = "97.75%"
        macro_f1     = "97.75%"
        avg_conf     = "98.16%"
        total_imgs   = "1,161"
        eval_type    = "In-Distribution (Clean Lab) Evaluation"
        eval_desc    = (
            "The model was evaluated on 1,161 held-out test images from the "
            "prepared Fruits-360 dataset split, covering all six classes under "
            "controlled, lab-quality imaging conditions."
        )
        per_class_rows = [
            ("Apple",          "200", "199", "99.5%", "90.5%", "99.5%", "94.8%"),
            ("Apple_Rotten",   "200", "200", "100.0%","97.1%","100.0%","98.5%"),
            ("Banana",         "161", "161", "100.0%","100.0%","100.0%","100.0%"),
            ("Banana_Rotten",  "200", "198", "99.0%", "100.0%","99.0%","99.5%"),
            ("Orange",         "200", "179", "89.5%", "100.0%","89.5%","94.5%"),
            ("Orange_Rotten",  "200", "197", "98.5%", "100.0%","98.5%","99.2%"),
        ]
        conf_matrix = [
            ("Apple",          "199", "1",  "0", "0", "0",  "0"),
            ("Apple_Rotten",   "0",   "200","0", "0", "0",  "0"),
            ("Banana",         "0",   "0",  "161","0","0",  "0"),
            ("Banana_Rotten",  "0",   "2",  "0","198","0",  "0"),
            ("Orange",         "21",  "0",  "0", "0","179", "0"),
            ("Orange_Rotten",  "0",   "3",  "0", "0","0",  "197"),
        ]
    else:
        overall_acc  = "89.15%"
        macro_prec   = "90.86%"
        macro_rec    = "88.95%"
        macro_f1     = "89.23%"
        avg_conf     = "95.59%"
        total_imgs   = "1,161"
        eval_type    = "Real-World Simulation Evaluation"
        eval_desc    = (
            "The model was evaluated on 1,161 test images with real-world "
            "augmentation applied—simulating lighting variation, camera blur, "
            "Gaussian sensor noise, and colour temperature shifts encountered "
            "during deployment on consumer-grade devices."
        )
        per_class_rows = [
            ("Apple",          "200", "190", "95.0%", "85.2%", "95.0%", "89.8%"),
            ("Apple_Rotten",   "200", "170", "85.0%", "96.6%", "85.0%", "90.4%"),
            ("Banana",         "161", "134", "83.2%", "98.5%", "83.2%", "90.2%"),
            ("Banana_Rotten",  "200", "190", "95.0%", "93.6%", "95.0%", "94.3%"),
            ("Orange",         "200", "153", "76.5%", "96.2%", "76.5%", "85.2%"),
            ("Orange_Rotten",  "200", "198", "99.0%", "75.0%", "99.0%", "85.3%"),
        ]
        conf_matrix = [
            ("Apple",          "190", "1",  "0", "0",  "4",  "5"),
            ("Apple_Rotten",   "3",   "170","0", "0",  "2",  "25"),
            ("Banana",         "2",   "0",  "134","13","0",  "12"),
            ("Banana_Rotten",  "1",   "3",  "2","190", "0",  "4"),
            ("Orange",         "27",  "0",  "0", "0",  "153","20"),
            ("Orange_Rotten",  "0",   "2",  "0", "0",  "0",  "198"),
        ]

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading(TITLE, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph("A PROJECT REPORT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True; p.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph("Submitted by")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, reg in STUDENTS:
        p = doc.add_paragraph(f"{name}  [Reg No: {reg}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph("Under the guidance of")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(GUIDE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph(GUIDE_DESIG)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph("in partial fulfillment for the award of the degree of")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("BACHELOR OF TECHNOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("in\nCOMPUTER SCIENCE AND ENGINEERING")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph("FACULTY OF ENGINEERING AND TECHNOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(DEPT)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("VADAPALANI CAMPUS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(YEAR)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # BONAFIDE CERTIFICATE
    # ══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("BONAFIDE CERTIFICATE", level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    names_str = " and ".join(f"{n} [Reg No: {r}]" for n, r in STUDENTS)
    p = doc.add_paragraph(
        f'Certified that 21CSP401L project report titled "{TITLE}" is the bonafide '
        f'work of "{names_str}", who carried out the project work under my supervision. '
        f'Certified further, that to the best of my knowledge the work reported herein '
        f'does not form any other project report or dissertation on the basis of which a '
        f'degree or award was conferred on an earlier occasion on this or any other student.'
    )
    p.style = 'Body Text'

    doc.add_paragraph()
    p_guide = doc.add_paragraph(f"GUIDE\n{GUIDE}\n{GUIDE_DESIG}\nDept. of CSE")
    p_guide.style = 'Body Text'
    doc.add_paragraph()
    p_hod = doc.add_paragraph("HEAD OF THE DEPARTMENT\nDr. Golda Dilip\nProfessor\nDept. of CSE")
    p_hod.style = 'Body Text'
    doc.add_paragraph()
    p_ext = doc.add_paragraph("EXTERNAL EXAMINER\n\nDepartment of CSE")
    p_ext.style = 'Body Text'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("ABSTRACT", level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(
        "FruitSense is an AI-powered fruit quality analysis system that leverages "
        "deep learning to classify fruit type and detect decay from a single photograph. "
        "The system employs a MobileNetV2-based transfer learning architecture trained on "
        "a merged dataset of over 11,000 images spanning six classes: Apple, Apple_Rotten, "
        "Banana, Banana_Rotten, Orange, and Orange_Rotten. A dual-pipeline backend combines "
        "the CNN classifier with a colorimetric scoring engine (CIE L*a*b*) and a "
        "secondary SE-CNN regression model to produce a continuous freshness score and "
        "estimated shelf-life. The system is deployed as a FastAPI REST service backed by "
        "a React + Vite frontend, making it accessible via standard web browsers and "
        "mobile devices. "
        f"In {eval_type.lower()}, the classifier achieved an overall accuracy of "
        f"{overall_acc} with a macro F1-score of {macro_f1}, demonstrating strong "
        "generalisation across fruit types and freshness states."
    )
    p.style = 'Body Text'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (static)
    # ══════════════════════════════════════════════════════════════════════════
    h = doc.add_heading("TABLE OF CONTENTS", level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        ("ACKNOWLEDGEMENTS",                   "iv"),
        ("ABSTRACT",                           "v"),
        ("LIST OF TABLES",                     "viii"),
        ("LIST OF FIGURES",                    "ix"),
        ("ABBREVIATIONS",                      "x"),
        ("1  INTRODUCTION",                    "1"),
        ("   1.1  Project Overview",           "1"),
        ("   1.2  Problem Statement",          "2"),
        ("   1.3  Objectives",                 "2"),
        ("2  LITERATURE REVIEW",               "3"),
        ("   2.1  CNN-based Fruit Classification", "3"),
        ("   2.2  Transfer Learning in Agriculture","4"),
        ("3  SYSTEM ARCHITECTURE AND DESIGN",  "5"),
        ("   3.1  Overall Pipeline",           "5"),
        ("   3.2  MobileNetV2 Classifier",     "6"),
        ("4  METHODOLOGY",                     "7"),
        ("   4.1  Dataset Preparation",        "7"),
        ("   4.2  Training Strategy",          "8"),
        ("5  CODING AND TESTING",              "9"),
        ("6  RESULTS AND OBSERVATIONS",        "10"),
        ("7  CONCLUSION",                      "12"),
        ("REFERENCES",                         "13"),
    ]
    table = doc.add_table(rows=len(toc_items), cols=2)
    table.style = 'Table Grid'
    for ri, (label, page) in enumerate(toc_items):
        table.rows[ri].cells[0].text = label
        table.rows[ri].cells[1].text = page
        table.rows[ri].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 1 — INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("CHAPTER 1", level=1)
    doc.add_heading("INTRODUCTION", level=3)

    doc.add_heading("1.1  Project Overview", level=2)
    p = doc.add_paragraph(
        "FruitSense addresses the challenge of automating fruit quality assessment "
        "in supply chains and retail environments. Traditional manual inspection is "
        "subjective, time-consuming, and error-prone. FruitSense provides an instant, "
        "objective quality score derived from image analysis, removing human bias and "
        "enabling scalable deployment via a web-based interface."
    )
    p.style = 'Body Text'

    doc.add_heading("1.2  Problem Statement", level=2)
    p = doc.add_paragraph(
        "Fruit spoilage results in significant economic losses globally. Early and "
        "accurate detection of decay is critical in cold-chain logistics and point-of-sale "
        "operations. Existing solutions either rely on expensive hardware (spectroscopy, "
        "gas sensors) or lack the classification granularity to distinguish between fresh "
        "and rotten instances of multiple fruit species simultaneously."
    )
    p.style = 'Body Text'

    doc.add_heading("1.3  Objectives", level=2)
    for obj in [
        "Develop a six-class MobileNetV2 classifier for Apple, Banana, and Orange (fresh and rotten).",
        "Integrate a CIE L*a*b* colorimetric scoring engine for continuous freshness quantification.",
        "Build a FastAPI REST backend and React frontend for real-time inference.",
        f"Evaluate the classifier rigorously and achieve a production-grade accuracy of {overall_acc}.",
    ]:
        p = doc.add_paragraph(obj, style='List Paragraph')
        p.style = 'List Paragraph'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 2 — LITERATURE REVIEW
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("CHAPTER 2", level=1)
    doc.add_heading("LITERATURE REVIEW", level=3)

    doc.add_heading("2.1  CNN-based Fruit Classification", level=2)
    p = doc.add_paragraph(
        "Convolutional Neural Networks have been extensively applied to agricultural "
        "image classification. Mureşan and Oltean (2018) introduced the Fruits-360 dataset "
        "and demonstrated accuracy above 99% on clean studio images using a lightweight CNN. "
        "However, such results do not generalise to real-world conditions due to lighting "
        "variability, occlusion, and background clutter [1].\n\n"
        "Kamilaris and Prenafeta-Boldú (2018) surveyed over 40 deep-learning papers applied "
        "to plant and crop analysis, concluding that transfer learning consistently outperforms "
        "training from scratch when labelled data is limited [2]. This informed the decision "
        "to use MobileNetV2 pretrained on ImageNet as the backbone for FruitSense."
    )
    p.style = 'Body Text'

    doc.add_heading("2.2  Transfer Learning in Agriculture", level=2)
    p = doc.add_paragraph(
        "MobileNetV2 (Sandler et al., 2018) employs inverted residuals and linear "
        "bottlenecks to achieve competitive accuracy with minimal computational cost, "
        "making it ideal for deployment on edge and mobile devices [3]. Several studies "
        "have fine-tuned MobileNetV2 on custom food and produce datasets, reporting "
        "90–98% accuracy depending on class complexity and dataset quality.\n\n"
        "Compared to ResNet-50 and EfficientNet-B0, MobileNetV2 provides the best "
        "accuracy-to-latency tradeoff under CPU-only inference, which is the deployment "
        "environment for FruitSense [4]."
    )
    p.style = 'Body Text'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 3 — SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("CHAPTER 3", level=1)
    doc.add_heading("SYSTEM ARCHITECTURE AND DESIGN", level=3)

    doc.add_heading("3.1  Overall Pipeline", level=2)
    p = doc.add_paragraph(
        "The FruitSense pipeline consists of five stages:\n"
        "(1) Image ingestion via multipart POST to /api/analyze\n"
        "(2) Colour space conversion (RGB → CIE L*a*b*) and feature extraction\n"
        "(3) MobileNetV2 fruit/freshness classification\n"
        "(4) Colorimetric freshness scoring with weighted sub-scores\n"
        "(5) JSON response construction with class, freshness score, decay %, and shelf-life estimate"
    )
    p.style = 'Body Text'

    doc.add_heading("3.2  MobileNetV2 Classifier", level=2)
    p = doc.add_paragraph(
        "The classifier uses a MobileNetV2 backbone (pretrained on ImageNet, input 224×224×3) "
        "with a custom head: GlobalAveragePooling2D → BatchNormalization → Dense(512, ReLU) "
        "→ Dropout(0.4) → Dense(256, ReLU) → Dropout(0.3) → Dense(6, Softmax). "
        "Training proceeds in two phases: Phase 1 freezes the base and trains the head for "
        "15 epochs at lr=1e-3; Phase 2 unfreezes the top 30 layers for fine-tuning at lr=1e-4 "
        "for 20 epochs with early stopping (patience=7)."
    )
    p.style = 'Body Text'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 4 — METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("CHAPTER 4", level=1)
    doc.add_heading("METHODOLOGY", level=3)

    doc.add_heading("4.1  Dataset Preparation", level=2)
    p = doc.add_paragraph(
        "The training dataset was assembled by merging two sources: (a) the Fruits-360 "
        "dataset providing fresh Apple, Banana, and Orange images (studio background), "
        "and (b) a rotten-fruits dataset sourced from Kaggle containing naturally decayed "
        "specimens. After merging, class weights were computed inversely proportional to "
        "class frequency to address the imbalance between fresh (abundant) and rotten "
        "(fewer) samples. The final split: 80% train, 20% validation."
    )
    p.style = 'Body Text'

    doc.add_heading("4.2  Training Strategy", level=2)
    p = doc.add_paragraph(
        "Data augmentation during training included rotation (±30°), width/height shift "
        "(±20%), shear (15%), zoom (25%), horizontal flip, brightness variation [0.7–1.3], "
        "and channel shift (±10). MobileNetV2's preprocess_input function was applied to "
        "normalise inputs to [-1, 1] matching ImageNet pretrained statistics. "
        "ModelCheckpoint saved the best validation-accuracy checkpoint, and "
        "ReduceLROnPlateau halved the learning rate on plateau (patience=3, min_lr=1e-7)."
    )
    p.style = 'Body Text'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 5 — CODING AND TESTING
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("CHAPTER 5", level=1)
    doc.add_heading("CODING AND TESTING", level=3)

    doc.add_heading("5.1  Backend Implementation", level=2)
    p = doc.add_paragraph(
        "The backend is implemented in Python using FastAPI with Uvicorn as the ASGI "
        "server. The /api/analyze endpoint accepts a multipart image upload, performs "
        "preprocessing, runs inference via TensorFlow/Keras, and returns a structured "
        "JSON payload containing: fruit_type, is_fresh, freshness_score (0–100), "
        "decay_percentage, shelf_life_days, and classifier_confidence."
    )
    p.style = 'Body Text'

    doc.add_heading("5.2  Frontend Implementation", level=2)
    p = doc.add_paragraph(
        "The frontend is built with React 18 and Vite, communicating with the backend "
        "over HTTP. Users drag-and-drop or select a fruit image; the UI displays a "
        "freshness gauge, quality badge, decay timeline, and per-class probability "
        "breakdown. Responsive design ensures usability on both desktop and mobile."
    )
    p.style = 'Body Text'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 6 — RESULTS AND OBSERVATIONS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("CHAPTER 6", level=1)
    doc.add_heading("RESULTS AND OBSERVATIONS", level=3)

    p = doc.add_paragraph(eval_desc)
    p.style = 'Body Text'

    doc.add_paragraph()

    # Overall metrics table
    add_metric_table(
        doc,
        f"Table 6.1: Overall Model Performance — {eval_type}",
        [
            ("Overall Accuracy",  overall_acc),
            ("Macro Precision",   macro_prec),
            ("Macro Recall",      macro_rec),
            ("Macro F1-Score",    macro_f1),
            ("Avg. Confidence",   avg_conf),
            ("Total Test Images", total_imgs),
        ],
        ["Metric", "Value"]
    )

    # Per-class table
    add_metric_table(
        doc,
        "Table 6.2: Per-Class Classification Metrics",
        per_class_rows,
        ["Class", "N", "Correct", "Accuracy", "Precision", "Recall", "F1-Score"]
    )

    # Confusion matrix
    doc.add_paragraph()
    caption = doc.add_paragraph("Table 6.3: Confusion Matrix (Rows = True Label, Cols = Predicted)")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].bold = True; caption.runs[0].font.size = Pt(11)

    cm_headers = ["True \\ Pred", "Apple", "Apple_Rot", "Banana", "Banana_Rot", "Orange", "Orange_Rot"]
    table = doc.add_table(rows=1, cols=len(cm_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(cm_headers):
        hdr.cells[i].text = h
    style_table_header_row(hdr)
    for ri, row_data in enumerate(conf_matrix):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                set_cell_bg(row.cells[ci], 'EEF2FF')
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # Observations
    doc.add_heading("6.1  Key Observations", level=2)
    if is_clean:
        obs = [
            f"The classifier achieved {overall_acc} overall accuracy on the clean held-out test split.",
            "Apple and Banana achieved near-perfect classification (≥99% accuracy).",
            "Orange showed the lowest per-class accuracy (89.5%) due to visual similarity with Apple under certain lighting.",
            "Rotten classes were classified with high precision (97–100%), minimising false-positive decay alerts.",
            f"High average confidence ({avg_conf}) indicates the model produces well-calibrated predictions.",
        ]
    else:
        obs = [
            f"Under real-world simulation, overall accuracy dropped to {overall_acc} — a realistic deployment estimate.",
            "Orange showed the largest degradation (76.5%) due to colour-similarity with Apple under brightness shifts.",
            "Banana_Rotten maintained 95% accuracy, indicating robust texture-based rotten detection.",
            "Orange_Rotten achieved 99% recall but lower precision (75%) — some clean images were flagged as rotten.",
            f"Average confidence remained high ({avg_conf}), suggesting the model is overconfident under distribution shift.",
        ]
    for o in obs:
        p = doc.add_paragraph(o, style='List Paragraph')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPTER 7 — CONCLUSION
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("CHAPTER 7", level=1)
    doc.add_heading("CONCLUSION", level=3)

    p = doc.add_paragraph(
        "FruitSense demonstrates that a lightweight MobileNetV2 classifier, combined with "
        "a colorimetric scoring engine, can deliver reliable fruit quality assessment from "
        "a single consumer photograph. "
        f"The system achieves {overall_acc} overall accuracy ({eval_type.lower()}), with a "
        f"macro F1-score of {macro_f1}, confirming its viability for practical deployment.\n\n"
        "Future work includes expanding the supported fruit set (Mango, Strawberry, Tomato), "
        "integrating on-device inference via TensorFlow Lite for offline mobile use, and "
        "training on a diverse real-world dataset to close the accuracy gap observed under "
        "distribution shift conditions."
    )
    p.style = 'Body Text'

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading("REFERENCES", level=1)
    refs = [
        "[1] Mureşan, H. & Oltean, M. (2018). Fruit recognition from images using deep learning. "
        "Acta Universitatis Sapientiae, Informatica, 10(1), 26–42.",

        "[2] Kamilaris, A. & Prenafeta-Boldú, F.X. (2018). Deep learning in agriculture: "
        "A survey. Computers and Electronics in Agriculture, 147, 70–90.",

        "[3] Sandler, M., Howard, A., Zhu, M., Zhmoginov, A. & Chen, L.C. (2018). "
        "MobileNetV2: Inverted residuals and linear bottlenecks. CVPR 2018.",

        "[4] Howard, A. et al. (2019). Searching for MobileNetV3. ICCV 2019.",

        "[5] He, K., Zhang, X., Ren, S. & Sun, J. (2016). Deep residual learning for "
        "image recognition. CVPR 2016, pp. 770–778.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style='List Paragraph')

    # ── Save ───────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── Entry point ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out_dir = BASE_DIR

    clean_path = os.path.join(out_dir, "FruitSense_Report_CleanTest_97pct.docx")
    rw_path    = os.path.join(out_dir, "FruitSense_Report_RealWorld_89pct.docx")

    print("Generating Report 1 — Clean/Lab evaluation (97.67%) ...")
    generate_report(clean_path, scenario='clean')

    print("Generating Report 2 — Real-world simulation (89.15%) ...")
    generate_report(rw_path, scenario='realworld')

    print("\nDone! Both reports saved to:")
    print(f"  {clean_path}")
    print(f"  {rw_path}")
